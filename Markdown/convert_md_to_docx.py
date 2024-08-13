import markdown2
from bs4 import BeautifulSoup
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.opc.constants import RELATIONSHIP_TYPE

def add_hyperlink(paragraph, url, text):
    """A function that places a hyperlink within a paragraph object."""
    part = paragraph.part
    r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'Hyperlink')
    rPr.append(rStyle)
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

def markdown_to_html(md_file):
    try:
        with open(md_file, 'r', encoding='utf-8') as f:
            markdown_text = f.read()
        print("Markdown Text:")
        print(markdown_text)  # Debug: Print the raw Markdown text
    except Exception as e:
        print(f"Error reading markdown file: {e}")
        return None

    try:
        html = markdown2.markdown(markdown_text)
        print("Generated HTML:")
        print(html)  # Debug: Print the generated HTML
        return html
    except Exception as e:
        print(f"Error converting markdown to HTML: {e}")
        return None

def html_to_docx(html_content, docx_file):
    try:
        soup = BeautifulSoup(html_content, 'html.parser')
        doc = Document()

        for element in soup.children:
            if element.name == 'h1':
                doc.add_heading(element.text, level=1)
            elif element.name == 'h2':
                doc.add_heading(element.text, level=2)
            elif element.name == 'h3':
                doc.add_heading(element.text, level=3)
            elif element.name == 'p':
                doc.add_paragraph(element.text)
            elif element.name == 'ul':
                for li in element.find_all('li'):
                    doc.add_paragraph(li.text, style='ListBullet')
            elif element.name == 'ol':
                for li in element.find_all('li'):
                    doc.add_paragraph(li.text, style='ListNumber')
            elif element.name == 'a':
                p = doc.add_paragraph()
                add_hyperlink(p, element['href'], element.text)

        doc.save(docx_file)
        print(f"DOCX file created successfully: {docx_file}")
    except Exception as e:
        print(f"Error converting HTML to DOCX: {e}")

def markdown_to_word(md_file, docx_file):
    html_content = markdown_to_html(md_file)
    if html_content:
        html_to_docx(html_content, docx_file)

# Example usage
markdown_to_word('test.md', 'example.docx')
