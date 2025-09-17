# -*- coding: utf-8 -*-
"""
Markdown -> DOCX (pure Python, no Pandoc). Handles headings, paragraphs, bold/italic,
inline code, fenced code blocks, blockquotes, lists (ul/ol), links, and tables.
Usage:
  python md2docx.py "input.md" ["output.docx"]

Requires:
  pip install markdown beautifulsoup4 python-docx
"""

import sys
from pathlib import Path
from markdown import markdown
from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# ---------------------- Helpers ----------------------
def ensure_code_style(doc: Document):
    """Create a simple 'Code Block' paragraph style if missing."""
    styles = doc.styles
    if 'Code Block' not in [s.name for s in styles]:
        s = styles.add_style('Code Block', 1)  # WD_STYLE_TYPE.PARAGRAPH = 1
        s.font.name = 'Consolas'
        s.font.size = Pt(9)
        s.paragraph_format.space_before = Pt(6)
        s.paragraph_format.space_after = Pt(6)
    if 'Code Inline' not in [s.name for s in styles]:
        cs = styles.add_style('Code Inline', 2)  # WD_STYLE_TYPE.CHARACTER = 2
        cs.font.name = 'Consolas'
        cs.font.size = Pt(10)

def add_hyperlink(paragraph, url, text):
    """Insert a clickable hyperlink run into a paragraph."""
    # Create w:hyperlink element
    part = paragraph.part
    r_id = part.relate_to(url, reltype=qn('http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink'), is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    # Create a w:r for the link text
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    # Style: blue + underline
    u = OxmlElement('w:u'); u.set(qn('w:val'), 'single'); rPr.append(u)
    color = OxmlElement('w:color'); color.set(qn('w:val'), '0000FF'); rPr.append(color)
    new_run.append(rPr)
    t = OxmlElement('w:t'); t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

def add_runs_from_inline_html(p, node):
    """Render inline HTML (<strong>, <em>, <code>, <a>, and text) into a paragraph."""
    for child in node.children if isinstance(node, Tag) else [node]:
        if isinstance(child, NavigableString):
            if str(child):
                p.add_run(str(child))
        elif isinstance(child, Tag):
            if child.name in ('strong', 'b'):
                run = p.add_run(child.get_text())
                run.bold = True
            elif child.name in ('em', 'i'):
                run = p.add_run(child.get_text())
                run.italic = True
            elif child.name == 'code':
                run = p.add_run(child.get_text())
                try:
                    run.style = 'Code Inline'
                except Exception:
                    pass
            elif child.name == 'a':
                href = child.get('href') or ''
                text = child.get_text() or href
                if href:
                    add_hyperlink(p, href, text)
                else:
                    p.add_run(text)
            else:
                # Recurse for nested inline tags (e.g., <strong><code>...</code></strong>)
                add_runs_from_inline_html(p, child)

def add_list(doc, ul_or_ol: Tag, level=0):
    """Render <ul>/<ol> lists with basic nesting."""
    ordered = ul_or_ol.name == 'ol'
    style = "List Number" if ordered else "List Bullet"
    for li in ul_or_ol.find_all('li', recursive=False):
        # Paragraph for list item
        p = doc.add_paragraph(style=style)
        if level:
            p.paragraph_format.left_indent = Pt(18 * level)
        # Render LI content
        # Split inline (text) vs nested lists
        children = list(li.children)
        # First, text/inline before any nested list
        text_container = BeautifulSoup('', 'html.parser').new_tag('span')
        for c in children:
            if isinstance(c, Tag) and c.name in ('ul', 'ol'):
                break
            text_container.append(c if isinstance(c, NavigableString) else c.extract())
        add_runs_from_inline_html(p, text_container)

        # Then handle nested lists if present
        for c in li.children:
            if isinstance(c, Tag) and c.name in ('ul', 'ol'):
                add_list(doc, c, level=level+1)

def add_table(doc: Document, table_tag: Tag):
    rows = table_tag.find_all('tr', recursive=False)
    if not rows:
        return
    # Determine columns from first row
    first_cells = rows[0].find_all(['td', 'th'], recursive=False)
    cols = max(1, len(first_cells))
    table = doc.add_table(rows=0, cols=cols)
    # Header if <th> present
    header_cells = [c for c in first_cells if c.name == 'th']
    row_idx = 0
    if header_cells:
        hdr = table.add_row().cells
        for i, cell in enumerate(first_cells):
            hdr[i].text = cell.get_text(strip=True)
        row_idx += 1
        data_rows = rows[1:]
    else:
        data_rows = rows
    # Data rows
    for tr in data_rows:
        tds = tr.find_all(['td', 'th'], recursive=False)
        row = table.add_row().cells
        for i in range(cols):
            row[i].text = (tds[i].get_text(strip=True) if i < len(tds) else '')

def add_blockquote(doc: Document, bq: Tag):
    p = doc.add_paragraph(bq.get_text('\n', strip=True))
    try:
        p.style = 'Intense Quote'
    except Exception:
        pass

def add_pre_code(doc: Document, pre: Tag):
    code = pre.code.get_text() if pre.find('code') else pre.get_text()
    p = doc.add_paragraph(code.replace('\r\n', '\n').replace('\r', '\n'))
    try:
        p.style = 'Code Block'
    except Exception:
        pass

def add_paragraph(doc: Document, tag: Tag):
    p = doc.add_paragraph()
    add_runs_from_inline_html(p, tag)

# ---------------------- Converter ----------------------
def md_to_docx(md_path: str, docx_path: str | None = None) -> str:
    src = Path(md_path).expanduser().resolve()
    if not src.is_file():
        raise FileNotFoundError(f"Markdown not found: {src}")
    dest = Path(docx_path).expanduser().resolve() if docx_path else src.with_suffix(".docx")

    md_text = src.read_text(encoding="utf-8")

    # Enable markdown extensions: tables, fenced code, sane lists, etc.
    html = markdown(
        md_text,
        extensions=[
            "extra",            # includes abbr, attr_list, def_list, fenced_code, footnotes, tables
            "sane_lists",
            "toc",
            "smarty",
            "admonition"
        ],
        output_format="html5",
    )

    soup = BeautifulSoup(html, "html.parser")
    doc = Document()
    ensure_code_style(doc)

    for element in soup.body.children if soup.body else soup.children:
        if isinstance(element, NavigableString):
            text = str(element).strip()
            if text:
                doc.add_paragraph(text)
            continue
        if not isinstance(element, Tag):
            continue

        name = element.name
        if name in ("h1", "h2", "h3", "h4", "h5", "h6"):
            level = int(name[1])
            doc.add_heading(element.get_text(), level=level if 1 <= level <= 4 else 4)
        elif name == "p":
            add_paragraph(doc, element)
        elif name in ("ul", "ol"):
            add_list(doc, element, level=0)
        elif name == "table":
            add_table(doc, element)
        elif name == "blockquote":
            add_blockquote(doc, element)
        elif name == "pre":
            add_pre_code(doc, element)
        elif name in ("hr",):
            doc.add_page_break()
        elif name in ("div", "section"):
            # Render contents recursively
            for child in element.children:
                if isinstance(child, Tag) and child.name in ("p","ul","ol","table","pre","blockquote") or (isinstance(child, Tag) and child.name.startswith("h")):
                    # Let main loop handle via a small re-parse
                    add_fragment(doc, child)
                elif isinstance(child, NavigableString) and child.strip():
                    doc.add_paragraph(child.strip())
        else:
            # Fallback: treat as paragraph
            add_paragraph(doc, element)

    dest.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(dest))
    size = dest.stat().st_size
    if size <= 0:
        raise RuntimeError(f"Empty DOCX written: {dest}")
    return str(dest)

def add_fragment(doc: Document, tag: Tag):
    """Render a single tag using the same rules as the top-level loop."""
    name = tag.name
    if name in ("h1","h2","h3","h4","h5","h6"):
        level = int(name[1])
        doc.add_heading(tag.get_text(), level=level if 1 <= level <= 4 else 4)
    elif name == "p":
        add_paragraph(doc, tag)
    elif name in ("ul","ol"):
        add_list(doc, tag, level=0)
    elif name == "table":
        add_table(doc, tag)
    elif name == "blockquote":
        add_blockquote(doc, tag)
    elif name == "pre":
        add_pre_code(doc, tag)
    else:
        add_paragraph(doc, tag)

# ---------------------- CLI ----------------------
if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python md2docx.py <input.md> [output.docx]")
        sys.exit(1)
    src = sys.argv[1]
    out = sys.argv[2] if len(sys.argv) > 2 else None
    try:
        path = md_to_docx(src, out)
        print(f"✅ Saved: {path}")
    except Exception as e:
        sys.stderr.write(f"❌ Error: {e}\n")
        sys.exit(2)